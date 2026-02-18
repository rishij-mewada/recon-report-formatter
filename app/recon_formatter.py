"""
Recon Analytics Document Formatter

Creates professional Word documents matching Recon Analytics brand standards.
"""

import os
import re
import zipfile
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# Brand Colors (RGB)
COLORS = {
    "navy": RGBColor(0x20, 0x38, 0x64),  # #203864
    "gray_text": RGBColor(0x59, 0x59, 0x59),  # #595959
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "black": RGBColor(0x00, 0x00, 0x00),
}

# Hex colors for cell shading
HEX_COLORS = {
    "navy": "203864",
    "light_gray_border": "CCCCCC",
    "green_highlight": "E2EFDA",
    "red_highlight": "FCE4D6",
    "gray_text": "595959",
}


class ReconDocumentFormatter:
    """Formatter for creating Recon Analytics branded Word documents."""

    def __init__(self):
        self._caption_counters = {"table": 0, "figure": 0, "chart": 0}
        self.doc = None

    def reset_caption_counters(self):
        """Reset all caption counters for a new document."""
        self._caption_counters = {"table": 0, "figure": 0, "chart": 0}

    def extract_logo(
        self, source_docx: str, output_path: str = "/tmp/recon_logo.png"
    ) -> Optional[str]:
        """Extract Recon logo from existing branded document."""
        if not os.path.exists(source_docx):
            return None
        try:
            with zipfile.ZipFile(source_docx, "r") as zf:
                for name in zf.namelist():
                    if "word/media/" in name and name.endswith(".png"):
                        with zf.open(name) as src:
                            data = src.read()
                            # Logo is ~16KB, 1379x128 pixels
                            if len(data) < 50000:
                                with open(output_path, "wb") as dst:
                                    dst.write(data)
                                return output_path
        except Exception:
            pass
        return None

    def set_cell_shading(self, cell, fill_color: str):
        """Apply background color to table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:shd"))
        if existing is not None:
            tcPr.remove(existing)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_color)
        tcPr.append(shd)

    def set_cell_borders(self, cell, color: str = "CCCCCC"):
        """Set all four borders on table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        existing = tcPr.find(qn("w:tcBorders"))
        if existing is not None:
            tcPr.remove(existing)
        tcBorders = OxmlElement("w:tcBorders")
        for name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), color)
            tcBorders.append(border)
        tcPr.append(tcBorders)

    def _set_table_row_cant_split(self, table):
        """Prevent table rows from splitting across pages."""
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()

            # Set cantSplit to prevent row from breaking across pages
            cant_split = OxmlElement("w:cantSplit")
            cant_split.set(qn("w:val"), "true")
            trPr.append(cant_split)

    def _set_repeat_header_row(self, table):
        """Set the first row to repeat as header on each page."""
        if len(table.rows) > 0:
            tr = table.rows[0]._tr
            trPr = tr.get_or_add_trPr()

            # Set tblHeader to repeat this row at top of each page
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            trPr.append(tbl_header)

    def add_bottom_border(self, paragraph, color: str = "595959"):
        """Add bottom border to paragraph (for H2 headings)."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")  # 1.5pt
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def setup_document(self) -> Document:
        """Create and configure a new Recon Analytics document."""
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.page_width = Twips(12240)
        section.page_height = Twips(15840)
        section.top_margin = Twips(950)
        section.bottom_margin = Twips(1440)
        section.left_margin = Twips(851)
        section.right_margin = Twips(900)
        section.header_distance = Twips(426)
        section.footer_distance = Twips(288)

        # Configure styles
        styles = doc.styles

        # Normal style - 11pt body text
        normal = styles["Normal"]
        normal.font.name = "Calibri Light"
        normal.font.size = Pt(11)
        normal.font.color.rgb = COLORS["gray_text"]
        normal.paragraph_format.space_after = Pt(12)

        # Heading 1 - 16pt ALL CAPS title
        h1 = styles["Heading 1"]
        h1.font.name = "Calibri Light"
        h1.font.size = Pt(16)
        h1.font.bold = False
        h1.font.all_caps = True
        h1.font.color.rgb = COLORS["gray_text"]

        # Heading 2 - 12pt Bold section with border
        h2 = styles["Heading 2"]
        h2.font.name = "Calibri Light"
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = COLORS["gray_text"]
        h2.paragraph_format.space_before = Pt(18)

        # Heading 3 - 12pt Bold subsection
        h3 = styles["Heading 3"]
        h3.font.name = "Calibri Light"
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.font.color.rgb = COLORS["gray_text"]
        h3.paragraph_format.space_before = Pt(18)
        h3.paragraph_format.space_after = Pt(6)

        # Heading 4 - 12pt Bold Italic minor heading (black)
        h4 = styles["Heading 4"]
        h4.font.name = "Calibri Light"
        h4.font.size = Pt(12)
        h4.font.bold = True
        h4.font.italic = True
        h4.font.color.rgb = COLORS["black"]

        self.doc = doc
        return doc

    def add_page_number_header(self):
        """Add right-aligned page number to header."""
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        # First paragraph - page number right-aligned
        para = header.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.space_after = Pt(0)

        # Create PAGE field
        run = para.add_run()

        # Field begin
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        # Field instruction
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE   \\* MERGEFORMAT "
        run._r.append(instr)

        # Field separator
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        # Placeholder text
        num_text = OxmlElement("w:t")
        num_text.text = "1"
        run._r.append(num_text)

        # Field end
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

        # Format the run
        run.font.name = "Calibri Light"
        run.font.size = Pt(11)
        run.font.color.rgb = COLORS["gray_text"]

        # Add empty second paragraph for spacing
        para2 = header.add_paragraph()
        para2.paragraph_format.space_after = Pt(0)

    def add_table_of_contents(self):
        """Add properly configured Table of Contents."""
        # TOC Heading
        toc_heading = self.doc.add_paragraph("Table of Contents")
        toc_heading.style = "Heading 1"

        # TOC Field
        para = self.doc.add_paragraph()
        run = para.add_run()

        # Field begin
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        # Field instruction
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instr)

        # Field separator
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        # Placeholder text
        placeholder = OxmlElement("w:t")
        placeholder.text = (
            "Right-click and select 'Update Field' to generate Table of Contents"
        )
        run._r.append(placeholder)

        # Field end
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

        # Page break after TOC
        self.doc.add_page_break()

    def add_section_heading(self, title: str):
        """Add H2 section with bottom border."""
        heading = self.doc.add_heading(title, level=2)
        self.add_bottom_border(heading)
        return heading

    def add_subsection(self, title: str):
        """Add H3 subsection."""
        return self.doc.add_heading(title, level=3)

    def add_minor_heading(self, title: str):
        """Add H4 minor heading (bold italic)."""
        return self.doc.add_heading(title, level=4)

    def add_caption(self, caption_type: str, description: str):
        """Add a consistent caption for tables, figures, or charts."""
        self._caption_counters[caption_type] += 1
        num = self._caption_counters[caption_type]

        type_label = caption_type.capitalize()
        caption_text = f"{type_label} {num}: {description}"

        para = self.doc.add_paragraph()
        run = para.add_run(caption_text)
        run.font.name = "Calibri Light"
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = COLORS["gray_text"]

        para.paragraph_format.space_after = Pt(6)

        return para

    def add_table(
        self,
        headers: list[str],
        data: list[list],
        col_widths: Optional[list[int]] = None,
        highlights: Optional[dict] = None,
        numeric_cols: Optional[list[int]] = None,
        caption: Optional[str] = None,
    ):
        """
        Add a formatted Recon Analytics table with optional caption.

        Args:
            headers: List of column headers
            data: List of rows (each row is a list)
            col_widths: List of widths in DXA (default: equal widths)
            highlights: Dict of {(row, col): 'positive'|'negative'}
            numeric_cols: List of column indices to right-align
            caption: Optional caption description
        """
        highlights = highlights or {}
        numeric_cols = numeric_cols or []

        # Add caption if provided
        if caption:
            self.add_caption("table", caption)

        if col_widths is None:
            col_widths = [9360 // len(headers)] * len(headers)

        table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Prevent table rows from breaking across pages
        self._set_table_row_cant_split(table)

        # Repeat header row if table spans multiple pages
        self._set_repeat_header_row(table)

        # Header row
        for idx, (cell, text) in enumerate(zip(table.rows[0].cells, headers)):
            self.set_cell_shading(cell, HEX_COLORS["navy"])
            self.set_cell_borders(cell, HEX_COLORS["navy"])

            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            para.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT
                if idx in numeric_cols
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = para.add_run(str(text))
            run.font.name = "Calibri Light"
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = COLORS["white"]

        # Data rows
        for row_idx, row_data in enumerate(data):
            for col_idx, (cell, value) in enumerate(
                zip(table.rows[row_idx + 1].cells, row_data)
            ):
                # Conditional highlighting
                key = (row_idx, col_idx)
                if key in highlights:
                    fill = (
                        HEX_COLORS["green_highlight"]
                        if highlights[key] == "positive"
                        else HEX_COLORS["red_highlight"]
                    )
                    self.set_cell_shading(cell, fill)

                self.set_cell_borders(cell, HEX_COLORS["light_gray_border"])

                cell.text = ""
                para = cell.paragraphs[0]
                para.paragraph_format.space_after = Pt(0)
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.RIGHT
                    if col_idx in numeric_cols
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                self._add_formatted_runs(para, str(value))

        self.doc.add_paragraph()  # Space after table
        return table

    def add_figure(self, image_path: str, description: str, width_inches: float = 6.0):
        """Add an image with a Figure caption."""
        self.add_caption("figure", description)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        self.doc.add_paragraph()

    def add_chart(self, image_path: str, description: str, width_inches: float = 6.0):
        """Add a chart image with a Chart caption."""
        self.add_caption("chart", description)

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        self.doc.add_paragraph()

    def _add_formatted_runs(
        self, para, text: str, base_bold: bool = False, base_italic: bool = False
    ):
        """Parse inline **bold** and *italic* markers and add formatted runs to a paragraph."""
        if ("**" in text or "*" in text) and not base_bold:
            tokens = re.split(r"(\*\*[^*]+?\*\*|\*[^*]+?\*)", text)
            for token in tokens:
                if not token:
                    continue
                if token.startswith("**") and token.endswith("**"):
                    run = para.add_run(token[2:-2])
                    run.font.bold = True
                    run.font.italic = base_italic
                elif token.startswith("*") and token.endswith("*"):
                    run = para.add_run(token[1:-1])
                    run.font.italic = True
                    run.font.bold = False
                else:
                    run = para.add_run(token)
                    run.font.bold = False
                    run.font.italic = base_italic
                run.font.name = "Calibri Light"
                run.font.size = Pt(11)
                run.font.color.rgb = COLORS["gray_text"]
        else:
            run = para.add_run(text)
            run.font.name = "Calibri Light"
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS["gray_text"]
            run.font.italic = base_italic
            run.font.bold = base_bold

    def add_paragraph(self, text: str, italic: bool = False, bold: bool = False):
        """Add a styled body paragraph with support for inline **bold** and *italic* markers."""
        para = self.doc.add_paragraph()
        self._add_formatted_runs(para, text, base_bold=bold, base_italic=italic)
        return para

    def add_title_block(
        self,
        title: str,
        subtitle: Optional[str] = None,
        author: Optional[str] = None,
        date: Optional[str] = None,
    ):
        """Add the document title block."""
        # Main title
        self.doc.add_heading(title, level=1)

        if subtitle:
            p = self.doc.add_paragraph()
            run = p.add_run(subtitle)
            run.font.italic = True
            run.font.name = "Calibri Light"
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS["gray_text"]

        if author:
            p = self.doc.add_paragraph()
            run = p.add_run(f"By: {author}")
            run.font.name = "Calibri Light"
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS["gray_text"]

        if date:
            p = self.doc.add_paragraph()
            run = p.add_run(date)
            run.font.name = "Calibri Light"
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS["gray_text"]

    def add_footer(self, logo_path: Optional[str] = None):
        """Add branded footer with full-width logo and URL."""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        # Paragraph 1: Anchored logo
        para1 = footer.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para1.paragraph_format.space_after = Pt(0)

        # Add right indent
        pPr = para1._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:right"), "-491")
        pPr.append(ind)

        if logo_path and os.path.exists(logo_path):
            run = para1.add_run()
            rId, _ = run.part.get_or_add_image(logo_path)

            anchor_xml = f"""
            <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                       xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
                       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
                       distT="0" distB="0" distL="114300" distR="114300"
                       simplePos="0" relativeHeight="251658240" behindDoc="1"
                       locked="0" layoutInCell="1" allowOverlap="1">
                <wp:simplePos x="0" y="0"/>
                <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
                <wp:positionV relativeFrom="page"><wp:posOffset>9335744</wp:posOffset></wp:positionV>
                <wp:extent cx="7776000" cy="721774"/>
                <wp:effectExtent l="0" t="0" r="0" b="0"/>
                <wp:wrapNone/>
                <wp:docPr id="2" name="Logo"/>
                <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
                <a:graphic>
                    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                        <pic:pic>
                            <pic:nvPicPr><pic:cNvPr id="2" name="Logo"/><pic:cNvPicPr/></pic:nvPicPr>
                            <pic:blipFill>
                                <a:blip r:embed="{rId}"/>
                                <a:stretch><a:fillRect/></a:stretch>
                            </pic:blipFill>
                            <pic:spPr>
                                <a:xfrm><a:off x="0" y="0"/><a:ext cx="7776000" cy="721774"/></a:xfrm>
                                <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                            </pic:spPr>
                        </pic:pic>
                    </a:graphicData>
                </a:graphic>
            </wp:anchor>
            """
            anchor = parse_xml(anchor_xml)
            drawing = OxmlElement("w:drawing")
            drawing.append(anchor)
            run._r.append(drawing)

        # Paragraph 2: Spacer
        para2 = footer.add_paragraph()
        para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para2.paragraph_format.space_after = Pt(0)
        para2.paragraph_format.space_before = Pt(48)
        pPr2 = para2._p.get_or_add_pPr()
        ind2 = OxmlElement("w:ind")
        ind2.set(qn("w:right"), "-491")
        pPr2.append(ind2)

        # Paragraph 3: URL
        para3 = footer.add_paragraph()
        para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para3.paragraph_format.space_after = Pt(0)
        para3.paragraph_format.space_before = Pt(0)
        pPr3 = para3._p.get_or_add_pPr()
        ind3 = OxmlElement("w:ind")
        ind3.set(qn("w:right"), "-491")
        pPr3.append(ind3)

        url_run = para3.add_run("www.reconanalytics.com ")
        url_run.font.name = "Calibri Light"
        url_run.font.size = Pt(8)
        url_run.font.bold = True
        url_run.font.color.rgb = COLORS["navy"]

    def save(self, output_path: str) -> str:
        """Save the document to the specified path."""
        self.doc.save(output_path)
        return output_path
